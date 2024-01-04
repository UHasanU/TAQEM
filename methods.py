from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK


def get_sorted_instructors(instructorsColumn):
    instructors = set()

    for instructor in instructorsColumn:
        if instructor.row == 1 or instructor.value == "مو موجود":
            continue
        instructors.add(instructor.value)

    return sorted(instructors)


def write_title_page(document, title, font, font_size, margins=0.5, picture=None):
    section = document.sections[0]
    section.top_margin = Inches(margins)
    section.bottom_margin = Inches(margins)
    section.left_margin = Inches(margins)
    section.right_margin = Inches(margins)

    title_para = document.add_paragraph()
    title = title_para.add_run(title)
    title.bold = True
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.font.name = font
    title.font.size = Pt(font_size)

    if picture is not None:
        pass  # Later

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def fix_imlaa(text):
    misspelled_words = ["مره", "مرا", "سهله", "اجوبه", "أجوبه", "فابده", "ساعه", "صراحه", "ماده", "دقيقه"]
    word_list = text.split()
    for index, word in enumerate(word_list):
        if word in misspelled_words or (word[:2] == "ال" and word[2:] in misspelled_words):
            word_list[index] = word[:-1] + "ة"
    return " ".join(word_list)


def set_cell_text(cell, text, alignment, bold=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        for run in paragraph.runs:
            run.font.bold = bold
            run.font.size = Pt(16)
            run.font.name = "Times New Roman"


def write_taqeem(document, instructor_name, sheet):
    inst_name = document.add_heading(instructor_name.title(), level=1)
    inst_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    inst_name.runs[0].bold = True
    inst_name.runs[0].font.name = 'Times New Roman'
    inst_name.runs[0].font.size = Pt(16)

    instructor_rows = [instructor.row for instructor in sheet['C'] if instructor.value == instructor_name]

    for row_number in instructor_rows:
        table = document.add_table(rows=8, cols=2)
        table.style = "Light List Accent 2"
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_cell = table.cell(0, 0).merge(table.cell(0, 1))
        set_cell_text(title_cell, instructor_name.title(), WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)

        for count, row in enumerate(table.rows[:-1]):
            columns = ['B', 'D', 'F', 'G', 'H', 'I', 'J']
            row.cells[0].width = Inches(5.25)
            row.cells[1].width = Inches(1.86)

            right_cell_value = str(sheet[f"{columns[count]}1"].value)
            left_cell_value = fix_imlaa(str(sheet[f"{columns[count]}{row_number}"].value))

            set_cell_text(table.cell(count + 1, 1), right_cell_value, WD_PARAGRAPH_ALIGNMENT.RIGHT)
            set_cell_text(table.cell(count + 1, 0), left_cell_value, WD_PARAGRAPH_ALIGNMENT.RIGHT)

        document.add_paragraph().add_run()
