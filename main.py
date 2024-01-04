from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from openpyxl import load_workbook
from methods import *

# Load the workbook
wb = load_workbook('PHYSICS.xlsx')

# Access a worksheet directly by name
ws = wb['ردود النموذج 1']

# Get the instructors from column C in the Excel file
instructors = get_sorted_instructors(ws['C'])

doc = Document()

write_title_page(doc, "تقييمات الفيزياء", "DecoType Naskh", 24)

for instructor in instructors:
    write_taqeem(doc, instructor, ws)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

doc.save("PHYSICS.docx")










